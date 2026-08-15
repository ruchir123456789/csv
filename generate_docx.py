import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_styled_document():
    doc = Document()
    
    # Set standard margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Styles helper
    def set_cell_background(cell, hex_color):
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)
        
    def add_styled_heading(text, level):
        h = doc.add_heading(text, level=level)
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(6)
        return h

    # Title
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run("CSV Analyser & Product Intelligence Platform")
    run_title.font.name = "Calibri"
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(16, 120, 90)

    # Subtitle
    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(18)
    run_sub = sub_p.add_run("Master Technical Architecture & Implementation Guide")
    run_sub.font.name = "Calibri"
    run_sub.font.size = Pt(14)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph(
        "A production-grade, full-stack AI catalog enrichment, web verification, and data analytics platform built with "
        "FastAPI, React (Vite), MongoDB Atlas, Open Icecat Live API, and DuckDuckGo Web Intelligence."
    )

    # Section 1: Executive Overview
    add_styled_heading("1. Executive Overview & System Architecture", level=1)
    doc.add_paragraph(
        "The CSV Analyser & Product Intelligence Platform transforms unstructured, raw, or incomplete consumer electronics "
        "and domestic appliance datasets into rich, structured product catalogs. The system performs automatic brand & model "
        "recognition, queries verified manufacturer datasheets via Open Icecat, uses real-time web scraping fallback for domestic "
        "appliance specs (inverter types, BEE star ratings, wattage, pricing), and cross-verifies each record with live search engine "
        "indexes to produce a quantified data Closeness Score (0.0% to 100.0%)."
    )

    # High level components bullet points
    doc.add_paragraph("Key Architectural Components:", style='List Bullet')
    p1 = doc.add_paragraph(style='List Bullet')
    p1.add_run("Frontend Layer: ").bold = True
    p1.add_run("React 18 + Vite with custom glassmorphic dark theme, interactive data tables, Chart.js analytics dashboard, and a live MongoDB Atlas connection badge.")

    p2 = doc.add_paragraph(style='List Bullet')
    p2.add_run("Backend Layer: ").bold = True
    p2.add_run("High-performance FastAPI (Python 3.10+) asynchronous service providing REST APIs, streaming multipart CSV ingestion, and automated OpenAPI Swagger docs.")

    p3 = doc.add_paragraph(style='List Bullet')
    p3.add_run("Catalog Intelligence Engines: ").bold = True
    p3.add_run("Open Icecat Live API connector + BeautifulSoup4 DuckDuckGo scraper with appliance knowledge base.")

    p4 = doc.add_paragraph(style='List Bullet')
    p4.add_run("Cloud Database: ").bold = True
    p4.add_run("MongoDB Atlas Cloud Database (Cluster0, csv_analyser_db) with Motor async driver and automated dataset history persistence.")

    p5 = doc.add_paragraph(style='List Bullet')
    p5.add_run("Deployment Infrastructure: ").bold = True
    p5.add_run("Single-service unified deployment on Render serving both FastAPI backend and React SPA.")

    # Section 2: Technology Stack
    add_styled_heading("2. Complete Technology Stack & Library Reference", level=1)

    add_styled_heading("2.1 Backend Libraries (Python 3.10+)", level=2)
    
    table_backend = doc.add_table(rows=1, cols=3)
    table_backend.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_backend.autofit = False

    headers_b = ["Library / Tool", "Version", "Purpose & Role in Project"]
    col_widths_b = [Inches(1.8), Inches(1.0), Inches(3.7)]

    hdr_cells_b = table_backend.rows[0].cells
    for i, title in enumerate(headers_b):
        hdr_cells_b[i].text = title
        hdr_cells_b[i].width = col_widths_b[i]
        set_cell_background(hdr_cells_b[i], "10B981")
        for p in hdr_cells_b[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)

    backend_data = [
        ("FastAPI", "^0.115.0", "Asynchronous REST API framework for defining endpoints, validation, and Swagger."),
        ("Uvicorn", "^0.32.0", "Lightning-fast ASGI web server hosting FastAPI locally and in production."),
        ("Motor", "^3.6.0", "Asynchronous, non-blocking MongoDB driver for Python built on asyncio."),
        ("PyMongo", "^4.9.0", "Core MongoDB Python driver handling BSON serialization and command execution."),
        ("dnspython", "^2.6.0", "Enables DNS SRV resolution required by MongoDB Atlas mongodb+srv:// URIs."),
        ("Pandas", "^2.2.0", "High-throughput data frame manipulation, header deduplication, and CSV export."),
        ("NumPy", "^1.26.0", "Numerical backend for Pandas operations and array processing."),
        ("HTTPX", "^0.28.0", "Asynchronous HTTP client for querying Open Icecat and search engine pages."),
        ("BeautifulSoup4", "^4.12.0", "HTML parser used to extract titles, snippets, specs, and price tags."),
        ("Pydantic", "^2.9.0", "Data validation, serialization, and typing schemas for API payloads."),
        ("Pydantic-Settings", "^2.15.0", "Structured environment settings loader reading from .env."),
        ("python-dotenv", "^1.0.1", "Loads local environment variables into system environment."),
        ("python-multipart", "^0.0.12", "Multipart form-data handler for streaming CSV file uploads.")
    ]

    for lib, ver, pur in backend_data:
        row_cells = table_backend.add_row().cells
        row_cells[0].text = lib
        row_cells[1].text = ver
        row_cells[2].text = pur
        for i in range(3):
            row_cells[i].width = col_widths_b[i]

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    add_styled_heading("2.2 Frontend Libraries (React 18 / 19 + Vite)", level=2)
    table_frontend = doc.add_table(rows=1, cols=3)
    table_frontend.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_frontend.autofit = False

    headers_f = ["Library / Package", "Version", "Purpose & Role in Project"]
    col_widths_f = [Inches(1.8), Inches(1.0), Inches(3.7)]

    hdr_cells_f = table_frontend.rows[0].cells
    for i, title in enumerate(headers_f):
        hdr_cells_f[i].text = title
        hdr_cells_f[i].width = col_widths_f[i]
        set_cell_background(hdr_cells_f[i], "06B6D4")
        for p in hdr_cells_f[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)

    frontend_data = [
        ("React & React-DOM", "^19.0.0", "Component-based reactive UI library managing state and DOM rendering."),
        ("Vite", "^6.0.0", "Fast build tool with HMR development server and optimized Rollup production bundler."),
        ("Lucide React", "^0.460.0", "Vector icon library providing UI symbols and indicators."),
        ("Chart.js", "^4.4.0", "HTML5 canvas charting engine for KPI graphs and analytics."),
        ("react-chartjs-2", "^5.2.0", "React component wrappers for Chart.js (Doughnut, Bar, Line)."),
        ("Axios", "^1.7.0", "Promise-based HTTP client for API calls with automatic base URL detection."),
        ("Vanilla CSS", "Custom", "Glassmorphic dark design system with CSS custom variables and animations.")
    ]

    for pkg, ver, pur in frontend_data:
        row_cells = table_frontend.add_row().cells
        row_cells[0].text = pkg
        row_cells[1].text = ver
        row_cells[2].text = pur
        for i in range(3):
            row_cells[i].width = col_widths_f[i]

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Section 3: Detailed Modules & Components
    add_styled_heading("3. Core Components & Modules Breakdown", level=1)

    modules = [
        ("3.1 Positional & Content Heuristics Parser (backend/app/services/enricher.py)",
         "Parses irregular, unstructured, or duplicated CSV columns (e.g. manufactu, manufactu, brand, description). "
         "Extracts brands, model identifiers, capacities, and variant descriptions using content recognition and dictionary lookup."),
        
        ("3.2 Open Icecat Live Connector (backend/app/services/icecat.py)",
         "Queries https://live.icecat.biz/api/ using shopname=openicecat-live to pull verified manufacturer datasheets, official product codes, high-resolution media URLs, and standardized technical specifications."),
        
        ("3.3 Web Intelligence & Scraping Fallback Engine (backend/app/services/enricher.py)",
         "For domestic/regional appliances not cataloged globally in Icecat, performs DuckDuckGo queries and HTML extraction for compressor types, inverter technology, 5-Star BEE energy ratings, and current market pricing."),
        
        ("3.4 DuckDuckGo Closeness Scoring Engine (backend/app/services/verifier.py)",
         "Audits every dataset row against real-time market data to compute a Closeness Score (0% to 100%), verifies brand and model match, and assigns an Overall Accuracy Grade (A+, A, B)."),
        
        ("3.5 MongoDB Atlas Cloud Persistence (backend/app/db/mongodb.py & dataset_repository.py)",
         "Connects to Cluster0 on MongoDB Atlas using Motor async driver with automated connection pooling, heartbeat ping probes, and history storage for the UI History Drawer."),
        
        ("3.6 Unified FastAPI Application (backend/main.py & backend/run.py)",
         "Configured with lifespan management for graceful database connections and automated static mounting of frontend/dist to serve the full React single-page app and REST APIs from a single server/port.")
    ]

    for title, desc in modules:
        add_styled_heading(title, level=2)
        doc.add_paragraph(desc)

    # Section 4: REST API Reference
    add_styled_heading("4. REST API Endpoint Reference", level=1)

    table_api = doc.add_table(rows=1, cols=3)
    table_api.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_api.autofit = False

    headers_a = ["Method & Route", "Description", "Payload / Parameters"]
    col_widths_a = [Inches(2.2), Inches(2.3), Inches(2.0)]

    hdr_cells_a = table_api.rows[0].cells
    for i, title in enumerate(headers_a):
        hdr_cells_a[i].text = title
        hdr_cells_a[i].width = col_widths_a[i]
        set_cell_background(hdr_cells_a[i], "3B82F6")
        for p in hdr_cells_a[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)

    api_endpoints = [
        ("GET /health", "Health check and live MongoDB Atlas status", "None"),
        ("GET /api/db/status", "Detailed Atlas connection diagnostics", "None"),
        ("POST /api/csv/enrich", "Ingests CSV, runs Icecat & Web enrichment", "file: UploadFile (multipart)"),
        ("POST /api/csv/verify", "Cross-verifies dataset against live search", "file: UploadFile (multipart)"),
        ("POST /api/csv/verify/single", "Instant single-product verification probe", "JSON: brand, model_code, desc"),
        ("GET /api/csv/datasets", "Retrieves past datasets from MongoDB Atlas", "limit: int (default 30)"),
        ("GET /api/csv/{file_id}/preview", "Returns paginated rows for dataset", "page: int, page_size: int"),
        ("GET /api/csv/{file_id}/summary", "Returns summary analytics and metrics", "file_id: string"),
        ("GET /api/csv/{file_id}/download-enriched", "Downloads enriched CSV with full specs", "file_id: string"),
        ("GET /api/csv/{file_id}/download-verified", "Downloads verified CSV with audit notes", "file_id: string"),
        ("DELETE /api/csv/{file_id}", "Deletes dataset from disk & MongoDB Atlas", "file_id: string"),
        ("GET /docs", "Interactive Swagger OpenAPI UI", "None")
    ]

    for route, desc, param in api_endpoints:
        row_cells = table_api.add_row().cells
        row_cells[0].text = route
        row_cells[1].text = desc
        row_cells[2].text = param
        for i in range(3):
            row_cells[i].width = col_widths_a[i]

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Section 5: Execution & Deployment
    add_styled_heading("5. Local Execution & Cloud Deployment Guide", level=1)

    add_styled_heading("5.1 Local Execution (1-Command Launcher)", level=2)
    doc.add_paragraph(
        "From the root folder (csv analyser/), run any of the following commands to start both Backend and Frontend:"
    )
    doc.add_paragraph("python start.py   (or: npm start, or double-click start.bat)", style='List Bullet')
    doc.add_paragraph("Frontend UI: http://localhost:5173  (or http://localhost:5174)", style='List Bullet')
    doc.add_paragraph("Backend API: http://127.0.0.1:8000  (Interactive Docs: http://127.0.0.1:8000/docs)", style='List Bullet')

    add_styled_heading("5.2 Render Cloud Deployment (1-Service Unified Setup)", level=2)
    doc.add_paragraph("To deploy on Render with a single Web Service:")
    doc.add_paragraph("1. Connect your GitHub repository (ruchir123456789/csv) to a Web Service on Render.", style='List Bullet')
    doc.add_paragraph("2. Leave 'Root Directory' blank/empty.", style='List Bullet')
    doc.add_paragraph("3. Build Command: npm --prefix frontend install && npm --prefix frontend run build && pip install -r backend/requirements.txt", style='List Bullet')
    doc.add_paragraph("4. Start Command: python backend/run.py", style='List Bullet')
    doc.add_paragraph("5. Add Environment Variables: MONGODB_URL, DATABASE_NAME=csv_analyser_db, PYTHON_VERSION=3.11.0", style='List Bullet')

    # Save to disk
    output_path = os.path.abspath("PROJECT_DOCUMENTATION.docx")
    doc.save(output_path)
    print(f"Successfully created: {output_path}")

if __name__ == "__main__":
    create_styled_document()
